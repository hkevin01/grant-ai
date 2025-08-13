"""
Enhanced AI-Powered Grant Writing Assistant for Grant-AI
Provides comprehensive writing assistance, document analysis, and proposal
generation
"""
import json
import logging
import os
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Optional

# Document processing
try:
    import docx
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# AI/ML imports
try:
    import openai
    OPENAI_AVAILABLE = True
except ImportError:
    OPENAI_AVAILABLE = False

try:
    import anthropic
    ANTHROPIC_AVAILABLE = True
except ImportError:
    ANTHROPIC_AVAILABLE = False


class AIWritingError(Exception):
    """Custom exception for AI writing assistant errors"""
    pass


class ProposalSection:
    """Represents a section of a grant proposal"""

    def __init__(self, title: str, content: str = "",
                 section_type: str = "content",
                 word_limit: Optional[int] = None,
                 requirements: List[str] = None):
        self.title = title
        self.content = content
        # content, narrative, budget, appendix
        self.section_type = section_type
        self.word_limit = word_limit
        self.requirements = requirements or []
        self.ai_suggestions = []
        self.review_feedback = []
        self.last_updated = datetime.now()

    def get_word_count(self) -> int:
        """Get current word count of the section"""
        return len(self.content.split())

    def is_within_limit(self) -> bool:
        """Check if content is within word limit"""
        if not self.word_limit:
            return True
        return self.get_word_count() <= self.word_limit

    def get_status(self) -> str:
        """Get section completion status"""
        if not self.content.strip():
            return "empty"
        elif not self.is_within_limit():
            return "over_limit"
        elif (self.get_word_count() <
              (self.word_limit * 0.5 if self.word_limit else 100)):
            return "needs_expansion"
        else:
            return "complete"


class GrantProposal:
    """Represents a complete grant proposal with multiple sections"""

    def __init__(self, title: str, grant_id: str, organization_id: str):
        self.title = title
        self.grant_id = grant_id
        self.organization_id = organization_id
        self.created_at = datetime.now()
        self.last_updated = datetime.now()
        self.sections: Dict[str, ProposalSection] = {}
        self.metadata = {
            "deadline": None,
            "funding_amount": None,
            "agency": None,
            "program": None,
            "submission_method": None
        }
        self.collaboration_log = []
        self.version_history = []

    def add_section(self, section: ProposalSection):
        """Add a section to the proposal"""
        self.sections[section.title] = section
        self.last_updated = datetime.now()

    def get_total_word_count(self) -> int:
        """Get total word count across all sections"""
        return sum(section.get_word_count()
                   for section in self.sections.values())

    def get_completion_status(self) -> Dict[str, Any]:
        """Get overall completion status"""
        total_sections = len(self.sections)
        if total_sections == 0:
            return {"status": "empty", "completion_percentage": 0}

        completed = sum(1 for section in self.sections.values()
                        if section.get_status() == "complete")

        status = ("in_progress" if completed < total_sections
                  else "complete")

        return {
            "status": status,
            "completion_percentage": (completed / total_sections) * 100,
            "completed_sections": completed,
            "total_sections": total_sections,
            "word_count": self.get_total_word_count()
        }


class AIGrantWritingAssistant:
    """Advanced AI-powered grant writing assistant"""

    def __init__(self):
        """Initialize the AI Grant Writing Assistant"""
        self.logger = logging.getLogger(__name__)
        self.ai_provider = self._detect_ai_provider()
        self.prompts = self._load_enhanced_prompts()
        self.templates = self._load_enhanced_templates()
        self.proposals: Dict[str, GrantProposal] = {}

    def _detect_ai_provider(self) -> str:
        """Detect which AI provider to use based on available API keys"""
        if OPENAI_AVAILABLE and os.getenv("OPENAI_API_KEY"):
            return "openai"
        elif ANTHROPIC_AVAILABLE and os.getenv("ANTHROPIC_API_KEY"):
            return "anthropic"
        else:
            return "fallback"

    def _load_enhanced_prompts(self) -> Dict[str, Any]:
        """Load enhanced prompts for AI-powered writing assistance"""
        return {
            "proposal_generation": {
                "specific_aims": (
                    "You are an expert grant writer helping to create "
                    "compelling Specific Aims for a research proposal.\n\n"
                    "Grant Details:\n"
                    "- Agency: {agency}\n"
                    "- Program: {program}\n"
                    "- Focus Area: {focus_area}\n"
                    "- Budget: {budget}\n\n"
                    "Organization Profile:\n{organization_profile}\n\n"
                    "Project Overview:\n{project_overview}\n\n"
                    "Please generate 2-3 specific aims that are:\n"
                    "1. Clear and concise\n"
                    "2. Feasible within the timeline and budget\n"
                    "3. Aligned with the agency's priorities\n"
                    "4. Innovative and impactful\n"
                    "5. Measurable with clear outcomes\n\n"
                    "Format each aim with:\n"
                    "- Aim [Number]: [Concise title]\n"
                    "- Objective: [What will be accomplished]\n"
                    "- Approach: [How it will be done]\n"
                    "- Expected Outcome: [What will be learned/achieved]\n\n"
                    "Make sure the aims build on each other logically and "
                    "address the overall project goals."
                ),

                "narrative": """You are an expert grant writer creating a compelling project narrative.

Grant Information:
- Title: {title}
- Agency: {agency}
- Focus Area: {focus_area}
- Target Audience: {target_audience}

Organization Details:
{organization_details}

Project Summary:
{project_summary}

Please write a compelling narrative that includes:

1. PROBLEM STATEMENT (2-3 paragraphs):
   - Clear articulation of the problem
   - Supporting evidence and statistics
   - Why this problem matters now

2. PROPOSED SOLUTION (3-4 paragraphs):
   - Your innovative approach
   - How it addresses the root causes
   - Why it will be effective

3. IMPACT & OUTCOMES (2-3 paragraphs):
   - Who will benefit and how
   - Measurable outcomes
   - Long-term impact

Use compelling language, include relevant data, and make a strong case for funding.""",

                "budget_narrative": """You are an expert at writing detailed budget justifications for grant proposals.

Project Details:
- Duration: {duration}
- Total Budget: {total_budget}
- Project Type: {project_type}

Budget Categories Requested:
{budget_categories}

Please write a detailed budget justification that:

1. Explains each cost category clearly
2. Demonstrates cost-effectiveness
3. Shows how expenses directly support project goals
4. Includes market research for major purchases
5. Addresses any potential reviewer concerns

For each category, provide:
- Clear explanation of need
- Detailed breakdown of costs
- Justification for amounts requested
- Connection to project activities

Make the justification compelling and thorough."""
            },

            "review_enhancement": {
                "clarity_check": """Review the following grant proposal text for clarity and readability:

{text}

Provide feedback on:
1. Sentence structure and flow
2. Technical jargon that needs simplification
3. Unclear or ambiguous statements
4. Suggestions for improved readability
5. Overall comprehension level

Rate the clarity on a scale of 1-10 and provide specific improvement suggestions.""",

                "alignment_analysis": """Analyze how well this proposal text aligns with the funder's priorities:

Funder: {funder}
Program: {program}
Funding Priorities: {priorities}

Proposal Text:
{text}

Provide analysis on:
1. Alignment with stated priorities (score 1-10)
2. Use of funder's preferred language/terminology
3. Missing connections to funder goals
4. Specific suggestions for better alignment
5. Key phrases/concepts to emphasize""",

                "competitiveness_analysis": """Assess the competitiveness of this proposal section:

Grant Program: {program}
Competition Level: {competition_level}
Success Rate: {success_rate}

Proposal Section:
{text}

Evaluate:
1. Strength of the argument (1-10)
2. Innovation and uniqueness
3. Feasibility and credibility
4. Impact potential
5. Areas for improvement to increase competitiveness

Provide specific recommendations to make this more competitive."""
            },

            "collaboration": {
                "writing_coach": """You are a grant writing coach providing guidance on proposal development.

Writer's Question/Challenge:
{question}

Current Draft (if any):
{draft}

Project Context:
{context}

Provide coaching advice that includes:
1. Direct answer to their question
2. Best practices for this type of content
3. Examples or templates if helpful
4. Common mistakes to avoid
5. Next steps for improvement

Be encouraging and constructive in your feedback.""",

                "peer_review": """Conduct a peer review of this grant proposal section:

Section Type: {section_type}
Word Limit: {word_limit}
Review Criteria: {criteria}

Content:
{content}

Provide a thorough peer review including:
1. Strengths of the current draft
2. Areas needing improvement
3. Missing elements
4. Clarity and flow issues
5. Specific, actionable recommendations
6. Overall assessment and score

Be constructive and specific in your feedback."""
            },

            "document_analysis": {
                "requirements_extraction": """Extract and analyze the requirements from this grant announcement:

Grant Announcement Text:
{announcement_text}

Extract and organize:
1. Application requirements and deadlines
2. Eligibility criteria
3. Evaluation criteria and weights
4. Required sections and page limits
5. Submission format requirements
6. Budget constraints
7. Special requirements or restrictions

Present in a clear, organized format that can guide proposal development.""",

                "compliance_check": """Check this proposal section for compliance with grant requirements:

Requirements:
{requirements}

Proposal Section:
{section_content}

Check for:
1. Adherence to format requirements
2. Inclusion of all required elements
3. Word/page limit compliance
4. Style and formatting guidelines
5. Missing required information

Provide a compliance checklist with pass/fail status for each requirement."""
            }
        }

    def _load_enhanced_templates(self) -> Dict[str, str]:
        """Load enhanced templates for different proposal sections"""
        return {
            "specific_aims_template": """SPECIFIC AIMS

[BRIEF PROJECT OVERVIEW - 1-2 sentences describing the overall project]

The long-term goal of this research is to [LONG-TERM OBJECTIVE]. The objective of this application is to [IMMEDIATE OBJECTIVE]. The central hypothesis is that [HYPOTHESIS STATEMENT]. This hypothesis has been formulated based on [PRELIMINARY DATA/OBSERVATIONS].

The rationale for the proposed research is that [RATIONALE/SIGNIFICANCE]. Upon completion of the proposed aims, we expect to [EXPECTED IMPACT].

Specific Aim 1: [AIM 1 TITLE]
[Detailed description of what will be accomplished, approach to be used, and expected outcome]

Specific Aim 2: [AIM 2 TITLE]
[Detailed description of what will be accomplished, approach to be used, and expected outcome]

Specific Aim 3: [AIM 3 TITLE] (if applicable)
[Detailed description of what will be accomplished, approach to be used, and expected outcome]

The proposed research is innovative because [INNOVATION STATEMENT]. The research is significant because [SIGNIFICANCE STATEMENT].""",

            "project_narrative_template": """PROJECT NARRATIVE

1. STATEMENT OF NEED
[Problem description with supporting data and evidence]

2. PROJECT DESCRIPTION
[Detailed description of proposed activities and approach]

3. GOALS AND OBJECTIVES
[Clear, measurable goals and specific objectives]

4. METHODOLOGY/APPROACH
[How the work will be accomplished, including timeline]

5. EVALUATION PLAN
[How success will be measured and assessed]

6. SUSTAINABILITY
[How the project will continue beyond the grant period]

7. ORGANIZATIONAL CAPACITY
[Why your organization is qualified to do this work]""",

            "budget_justification_template": """BUDGET JUSTIFICATION

PERSONNEL
Principal Investigator: [NAME] - [% EFFORT]
[Justification for PI time and effort]

Research Staff: [NAMES AND ROLES]
[Justification for each position and level of effort]

EQUIPMENT
[List major equipment items with justification for need]

SUPPLIES
[Categories of supplies with estimated costs and justification]

TRAVEL
[Travel requirements with destinations and justification]

OTHER DIRECT COSTS
[Any other project-related expenses with justification]

INDIRECT COSTS
[Institution's approved indirect cost rate and calculation]

TOTAL PROJECT COST: $[AMOUNT]"""
        }

    async def generate_section_content(self, section_type: str, context: Dict[str, Any]) -> Dict[str, Any]:
        """Generate content for a specific proposal section using AI"""

        if section_type not in self.prompts["proposal_generation"]:
            return {"error": f"Unknown section type: {section_type}"}

        prompt = self.prompts["proposal_generation"][section_type].format(**context)

        try:
            if self.ai_provider == "openai":
                content = await self._generate_with_openai(prompt)
            elif self.ai_provider == "anthropic":
                content = await self._generate_with_anthropic(prompt)
            else:
                content = self._generate_fallback_content(section_type, context)

            return {
                "content": content,
                "word_count": len(content.split()),
                "generated_at": datetime.now().isoformat(),
                "ai_provider": self.ai_provider
            }

        except Exception as e:
            self.logger.error(f"Error generating content: {str(e)}")
            return {"error": f"Failed to generate content: {str(e)}"}

    async def _generate_with_openai(self, prompt: str) -> str:
        """Generate content using OpenAI API"""
        if not OPENAI_AVAILABLE:
            raise AIWritingError("OpenAI not available")

        client = openai.AsyncOpenAI()
        response = await client.chat.completions.create(
            model="gpt-4",
            messages=[
                {"role": "system", "content": "You are an expert grant writer with extensive experience in successful proposal development."},
                {"role": "user", "content": prompt}
            ],
            max_tokens=2000,
            temperature=0.7
        )

        return response.choices[0].message.content

    async def _generate_with_anthropic(self, prompt: str) -> str:
        """Generate content using Anthropic Claude API"""
        if not ANTHROPIC_AVAILABLE:
            raise AIWritingError("Anthropic not available")

        client = anthropic.AsyncAnthropic()
        response = await client.messages.create(
            model="claude-3-sonnet-20240229",
            max_tokens=2000,
            messages=[
                {"role": "user", "content": f"You are an expert grant writer. {prompt}"}
            ]
        )

        return response.content[0].text

    def _generate_fallback_content(self, section_type: str, context: Dict[str, Any]) -> str:
        """Generate fallback content when AI APIs are not available"""
        templates = {
            "specific_aims": self.templates.get("specific_aims_template", ""),
            "narrative": self.templates.get("project_narrative_template", ""),
            "budget_narrative": self.templates.get("budget_justification_template", "")
        }

        template = templates.get(section_type, "")

        # Simple placeholder replacement
        for key, value in context.items():
            placeholder = f"{{{key.upper()}}}"
            if placeholder in template:
                template = template.replace(placeholder, str(value))

        return template + "\n\n[Note: This is a template. AI assistance not available - please fill in the placeholders with your specific content.]"

    async def review_section(self, content: str, review_type: str, context: Dict[str, Any]) -> Dict[str, Any]:
        """Review a proposal section using AI"""

        if review_type not in self.prompts["review_enhancement"]:
            return {"error": f"Unknown review type: {review_type}"}

        context["text"] = content
        prompt = self.prompts["review_enhancement"][review_type].format(**context)

        try:
            if self.ai_provider == "openai":
                feedback = await self._generate_with_openai(prompt)
            elif self.ai_provider == "anthropic":
                feedback = await self._generate_with_anthropic(prompt)
            else:
                feedback = self._generate_fallback_review(content, review_type)

            return {
                "feedback": feedback,
                "review_type": review_type,
                "reviewed_at": datetime.now().isoformat(),
                "word_count": len(content.split()),
                "ai_provider": self.ai_provider
            }

        except Exception as e:
            self.logger.error(f"Error reviewing content: {str(e)}")
            return {"error": f"Failed to review content: {str(e)}"}

    def _generate_fallback_review(self, content: str, review_type: str) -> str:
        """Generate fallback review when AI APIs are not available"""
        word_count = len(content.split())
        paragraph_count = len(content.split('\n\n'))

        basic_feedback = f"""BASIC REVIEW FEEDBACK ({review_type.upper()})

Content Statistics:
- Word count: {word_count}
- Character count: {len(content)}
- Paragraphs: {paragraph_count}

Basic Recommendations:
1. Review for clarity and conciseness
2. Ensure all key points are covered
3. Check for proper grammar and spelling
4. Verify alignment with grant requirements
5. Consider adding more specific examples

Note: AI-powered detailed review not available. Please use human reviewers for
comprehensive feedback."""

        return basic_feedback

    def create_proposal(self, title: str, grant_id: str, organization_id: str) -> str:
        """Create a new grant proposal"""
        proposal_id = f"{organization_id}_{grant_id}_{datetime.now().strftime('%Y%m%d_%H%M%S')}"

        proposal = GrantProposal(title, grant_id, organization_id)
        self.proposals[proposal_id] = proposal

        return proposal_id

    def add_collaboration_note(self, proposal_id: str, author: str, note: str):
        """Add a collaboration note to a proposal"""
        if proposal_id in self.proposals:
            self.proposals[proposal_id].collaboration_log.append({
                "timestamp": datetime.now().isoformat(),
                "author": author,
                "note": note
            })

    def export_proposal_to_docx(self, proposal_id: str, output_path: str) -> bool:
        """Export proposal to Word document"""
        if not DOCX_AVAILABLE:
            self.logger.error("python-docx not available for Word export")
            return False

        if proposal_id not in self.proposals:
            return False

        proposal = self.proposals[proposal_id]

        try:
            doc = docx.Document()

            # Add title
            doc.add_heading(proposal.title, 0)

            # Add metadata
            doc.add_paragraph(f"Grant ID: {proposal.grant_id}")
            doc.add_paragraph(f"Organization: {proposal.organization_id}")
            doc.add_paragraph(f"Created: {proposal.created_at.strftime('%Y-%m-%d')}")
            doc.add_paragraph("")

            # Add sections
            for section_title, section in proposal.sections.items():
                doc.add_heading(section_title, 1)
                doc.add_paragraph(section.content)
                doc.add_paragraph("")

            doc.save(output_path)
            return True

        except Exception as e:
            self.logger.error(f"Error exporting to Word: {str(e)}")
            return False

    def import_proposal_from_docx(self, file_path: str) -> Optional[str]:
        """Import proposal from Word document"""
        if not DOCX_AVAILABLE:
            self.logger.error("python-docx not available for Word import")
            return None

        try:
            doc = docx.Document(file_path)

            # Extract text content
            content = []
            current_section = None

            for paragraph in doc.paragraphs:
                if paragraph.style.name.startswith('Heading'):
                    if current_section:
                        content.append(current_section)
                    current_section = {
                        "title": paragraph.text,
                        "content": ""
                    }
                elif current_section:
                    current_section["content"] += paragraph.text + "\n"

            if current_section:
                content.append(current_section)

            # Create proposal
            file_name = Path(file_path).stem
            proposal_id = self.create_proposal(
                title=file_name,
                grant_id="imported",
                organization_id="imported"
            )

            # Add sections
            proposal = self.proposals[proposal_id]
            for section_data in content:
                section = ProposalSection(
                    title=section_data["title"],
                    content=section_data["content"]
                )
                proposal.add_section(section)

            return proposal_id

        except Exception as e:
            self.logger.error(f"Error importing from Word: {str(e)}")
            return None

    def get_proposal_status(self, proposal_id: str) -> Optional[Dict[str, Any]]:
        """Get detailed status of a proposal"""
        if proposal_id not in self.proposals:
            return None

        proposal = self.proposals[proposal_id]
        status = proposal.get_completion_status()

        section_details = {}
        for title, section in proposal.sections.items():
            section_details[title] = {
                "word_count": section.get_word_count(),
                "word_limit": section.word_limit,
                "status": section.get_status(),
                "within_limit": section.is_within_limit(),
                "requirements_met": len(section.requirements),
                "ai_suggestions": len(section.ai_suggestions),
                "last_updated": section.last_updated.isoformat()
            }

        return {
            "proposal_info": {
                "title": proposal.title,
                "grant_id": proposal.grant_id,
                "organization_id": proposal.organization_id,
                "created_at": proposal.created_at.isoformat(),
                "last_updated": proposal.last_updated.isoformat()
            },
            "completion_status": status,
            "sections": section_details,
            "collaboration": {
                "notes_count": len(proposal.collaboration_log),
                "versions_count": len(proposal.version_history)
            }
        }

    def save_proposals(self, file_path: str):
        """Save all proposals to file"""
        try:
            proposals_data = {}
            for proposal_id, proposal in self.proposals.items():
                sections_data = {}
                for section_title, section in proposal.sections.items():
                    sections_data[section_title] = {
                        "title": section.title,
                        "content": section.content,
                        "section_type": section.section_type,
                        "word_limit": section.word_limit,
                        "requirements": section.requirements,
                        "ai_suggestions": section.ai_suggestions,
                        "review_feedback": section.review_feedback,
                        "last_updated": section.last_updated.isoformat()
                    }

                proposals_data[proposal_id] = {
                    "title": proposal.title,
                    "grant_id": proposal.grant_id,
                    "organization_id": proposal.organization_id,
                    "created_at": proposal.created_at.isoformat(),
                    "last_updated": proposal.last_updated.isoformat(),
                    "sections": sections_data,
                    "metadata": proposal.metadata,
                    "collaboration_log": proposal.collaboration_log,
                    "version_history": proposal.version_history
                }

            with open(file_path, 'w') as f:
                json.dump(proposals_data, f, indent=2)

        except Exception as e:
            self.logger.error(f"Error saving proposals: {str(e)}")

    def load_proposals(self, file_path: str):
        """Load proposals from file"""
        try:
            if not Path(file_path).exists():
                return

            with open(file_path, 'r') as f:
                proposals_data = json.load(f)

            for proposal_id, proposal_data in proposals_data.items():
                proposal = GrantProposal(
                    title=proposal_data["title"],
                    grant_id=proposal_data["grant_id"],
                    organization_id=proposal_data["organization_id"]
                )
                proposal.created_at = datetime.fromisoformat(proposal_data["created_at"])
                proposal.last_updated = datetime.fromisoformat(proposal_data["last_updated"])
                proposal.metadata = proposal_data.get("metadata", {})
                proposal.collaboration_log = proposal_data.get("collaboration_log", [])
                proposal.version_history = proposal_data.get("version_history", [])

                # Load sections
                for section_title, section_data in proposal_data["sections"].items():
                    section = ProposalSection(
                        title=section_data["title"],
                        content=section_data["content"],
                        section_type=section_data.get("section_type", "content"),
                        word_limit=section_data.get("word_limit"),
                        requirements=section_data.get("requirements", [])
                    )
                    section.ai_suggestions = section_data.get("ai_suggestions", [])
                    section.review_feedback = section_data.get("review_feedback", [])
                    section.last_updated = datetime.fromisoformat(section_data["last_updated"])

                    proposal.add_section(section)

                self.proposals[proposal_id] = proposal

        except Exception as e:
            self.logger.error(f"Error loading proposals: {str(e)}")


# Create global instance
ai_writing_assistant = AIGrantWritingAssistant()
